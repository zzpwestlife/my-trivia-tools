import os
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_md_to_docx(md_file_path, output_docx_path=None):
    if not output_docx_path:
        output_docx_path = os.path.splitext(md_file_path)[0] + '.docx'

    # Read Markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert to HTML
    html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create Document
    document = Document()
    
    # Base directory for resolving relative paths
    base_dir = os.path.dirname(md_file_path)

    # Helper to process children (handles inline formatting like bold/italic)
    def process_element_children(element, paragraph):
        for child in element.children:
            if child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.color.rgb = RGBColor(200, 50, 50)  # Reddish for inline code
            elif child.name == 'a':
                # Simplified link handling: text (url)
                text = child.get_text()
                href = child.get('href', '')
                run = paragraph.add_run(f"{text} ({href})")
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            elif child.name == 'img':
                # Inline image? Add as block for simplicity
                src = child.get('src')
                alt = child.get('alt', '')
                add_image(src, alt)
            elif child.name is None:
                # NavigableString
                paragraph.add_run(child.string)
            else:
                # Recurse for nested tags (like bold inside italic)
                process_element_children(child, paragraph)

    def add_image(src, alt):
        full_path = os.path.join(base_dir, src)
        # Remove query parameters
        if '?' in full_path:
            full_path = full_path.split('?')[0]
            
        if os.path.exists(full_path):
            try:
                document.add_picture(full_path, width=Inches(6.0))
                if alt:
                    caption = document.add_paragraph(alt)
                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption.style = 'Caption'
            except Exception as e:
                p = document.add_paragraph()
                run = p.add_run(f"[Image upload failed: {src}]")
                run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            p = document.add_paragraph()
            run = p.add_run(f"[Image not found: {src}]")
            run.font.color.rgb = RGBColor(255, 0, 0)

    # Iterate over top-level elements
    for element in soup.body.children if soup.body else soup.children:
        if element.name == 'h1':
            document.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            document.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            document.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            document.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            document.add_heading(element.get_text(), level=5)
        elif element.name == 'h6':
            document.add_heading(element.get_text(), level=6)
            
        elif element.name == 'p':
            # Check if paragraph contains only an image
            img = element.find('img')
            if img and len(element.get_text(strip=True)) == 0:
                src = img.get('src')
                alt = img.get('alt', '')
                add_image(src, alt)
            else:
                p = document.add_paragraph()
                process_element_children(element, p)
                
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = document.add_paragraph(style='List Bullet')
                process_element_children(li, p)
                
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = document.add_paragraph(style='List Number')
                process_element_children(li, p)
                
        elif element.name == 'blockquote':
            p = document.add_paragraph(style='Quote')
            process_element_children(element, p)
            
        elif element.name == 'pre':
            # Code block
            code = element.find('code')
            text = code.get_text() if code else element.get_text()
            p = document.add_paragraph()
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            # Add border or background (requires intricate XML manipulation in python-docx, skipping for now)
            
        elif element.name == 'hr':
            document.add_paragraph('___________________________________________________')

    document.save(output_docx_path)
    print(f"Successfully converted to: {output_docx_path}")
    return output_docx_path

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        md_path = sys.argv[1]
        convert_md_to_docx(md_path)
    else:
        print("Usage: python src/md_converter.py <path_to_md_file>")
